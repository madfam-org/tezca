"""
Export views for downloading laws in multiple formats.

Tier system:
  - anon:    TXT only, 10/hour by IP
  - free:    TXT + PDF, 30/hour by user_id
  - premium: All 6 formats, 100/hour by user_id

Formats: txt, pdf, latex, docx, epub, json
"""

import io
import json as json_module
import logging
import re
from datetime import datetime

from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.template.loader import render_to_string
from drf_spectacular.utils import extend_schema
from rest_framework.decorators import api_view
from rest_framework.response import Response

from .config import INDEX_NAME, es_client
from .export_throttles import TIER_LIMITS, check_export_quota, log_export
from .models import Law

logger = logging.getLogger(__name__)

# Optional dependencies
try:
    from weasyprint import HTML as WeasyHTML

    _has_weasyprint = True
except (ImportError, OSError):
    _has_weasyprint = False

try:
    import jinja2

    _has_jinja2 = True
except (ImportError, OSError):
    _has_jinja2 = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt

    _has_docx = True
except (ImportError, OSError):
    _has_docx = False

try:
    from ebooklib import epub

    _has_ebooklib = True
except (ImportError, OSError):
    _has_ebooklib = False

# ── Format → minimum tier mapping ──────────────────────────────────────

FORMAT_TIERS = {
    "txt": "anon",
    "pdf": "free",
    "latex": "premium",
    "docx": "premium",
    "epub": "premium",
    "json": "premium",
}

TIER_RANK = {"anon": 0, "free": 1, "premium": 2}


# ── Helpers ────────────────────────────────────────────────────────────


def _get_user_tier(request) -> tuple[str, str]:
    """
    Determine user tier and user_id from request.user (set by CombinedAuthentication).

    Returns (tier, user_id). Falls back to ("anon", "").
    """
    user = getattr(request, "user", None)
    if user and getattr(user, "is_authenticated", False):
        tier = getattr(user, "tier", "free")
        # Map API key tiers to export tiers
        tier_map = {"internal": "premium", "pro": "premium", "enterprise": "premium"}
        tier = tier_map.get(tier, tier)
        if tier not in TIER_RANK:
            tier = "free"
        return (tier, getattr(user, "id", ""))
    return ("anon", "")


def _get_client_ip(request) -> str:
    """Extract client IP, respecting X-Forwarded-For."""
    xff = request.META.get("HTTP_X_FORWARDED_FOR")
    if xff:
        return xff.split(",")[0].strip()
    return request.META.get("REMOTE_ADDR", "127.0.0.1")


def _check_access(request, fmt: str):
    """
    Check tier access and quota for an export request.

    Returns (tier, user_id, ip) on success.
    Raises an HttpResponse (403/429) on failure via return.
    Returns tuple of 4 elements: (tier, user_id, ip, error_response_or_None).
    """
    tier, user_id = _get_user_tier(request)
    ip = _get_client_ip(request)
    required_tier = FORMAT_TIERS.get(fmt, "premium")

    # Tier check
    if TIER_RANK.get(tier, 0) < TIER_RANK.get(required_tier, 0):
        if tier == "anon":
            return (
                tier,
                user_id,
                ip,
                Response(
                    {
                        "error": "Authentication required",
                        "required_tier": required_tier,
                    },
                    status=403,
                ),
            )
        return (
            tier,
            user_id,
            ip,
            Response(
                {
                    "error": "Insufficient tier",
                    "your_tier": tier,
                    "required_tier": required_tier,
                },
                status=403,
            ),
        )

    # Quota check
    allowed, retry_after = check_export_quota(tier, user_id, ip)
    if not allowed:
        resp = Response(
            {
                "error": "Rate limit exceeded",
                "retry_after": retry_after,
                "limit": TIER_LIMITS.get(tier, 10),
            },
            status=429,
        )
        resp["Retry-After"] = str(retry_after)
        return tier, user_id, ip, resp

    return tier, user_id, ip, None


def _get_articles(law_id: str, max_articles: int = 10000) -> list[dict]:
    """Fetch all articles for a law from Elasticsearch."""
    try:
        es = es_client
        if not es.ping():
            return []

        result = es.search(
            index=INDEX_NAME,
            query={"match_phrase": {"law_id": law_id}},
            size=max_articles,
            sort=[{"article": {"order": "asc"}}],
            source=["article", "text"],
        )
        return [
            {
                "article_id": hit["_source"].get("article", ""),
                "text": hit["_source"].get("text", ""),
            }
            for hit in result["hits"]["hits"]
        ]
    except Exception:
        logger.warning("ES unavailable for export %s", law_id, exc_info=True)
        return []


def _tier_label(tier: str) -> str:
    return {"federal": "Federal", "state": "Estatal", "municipal": "Municipal"}.get(
        tier, tier.title()
    )


def _safe_filename(law_id: str) -> str:
    return law_id.replace("/", "_").replace(" ", "_")


def _law_context(law: Law, articles: list[dict]) -> dict:
    """Build common context dict for templates."""
    latest_version = law.versions.order_by("-publication_date").first()
    pub_date = None
    if latest_version and latest_version.publication_date:
        pub_date = latest_version.publication_date.strftime("%d de %B de %Y")

    state = None
    if law.tier == "state" and "_" in law.official_id:
        state = law.official_id.split("_")[0].replace("_", " ").title()

    return {
        "law_name": law.name,
        "official_id": law.official_id,
        "tier_label": _tier_label(law.tier),
        "category": law.category,
        "state": state,
        "status": law.status,
        "publication_date": pub_date,
        "article_count": len(articles),
        "articles": articles,
        "generation_date": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }


# ── Export views ───────────────────────────────────────────────────────


@extend_schema(
    tags=["Export"],
    summary="Export law as plain text",
    description="Download a law's full text as a UTF-8 plain text file. Available to all users.",
)
@api_view(["GET"])
def export_txt(request, law_id):
    """Export a law as clean formatted plain text."""
    tier, user_id, ip, error = _check_access(request, "txt")
    if error:
        return error

    law = get_object_or_404(Law, official_id=law_id)
    articles = _get_articles(law_id)

    if not articles:
        return Response({"error": "No articles found for this law."}, status=404)

    lines = []
    lines.append("=" * 72)
    lines.append(law.name.center(72))
    lines.append("=" * 72)
    lines.append("")

    tier_label = _tier_label(law.tier)
    lines.append(f"Tipo: {tier_label}")
    if law.category:
        lines.append(f"Categoría: {law.category}")
    if law.status:
        lines.append(f"Estado: {law.status}")

    latest_version = law.versions.order_by("-publication_date").first()
    if latest_version and latest_version.publication_date:
        lines.append(f"Publicado: {latest_version.publication_date}")

    lines.append(f"Artículos: {len(articles)}")
    lines.append("")
    lines.append("-" * 72)
    lines.append("")

    for article in articles:
        lines.append(article["article_id"])
        lines.append("")
        lines.append(article["text"])
        lines.append("")
        lines.append("")

    lines.append("-" * 72)
    lines.append(
        f"Generado por Tezca — El Espejo de la Ley | {datetime.now().strftime('%Y-%m-%d')} | tezca.mx"
    )
    lines.append("")

    content = "\n".join(lines)
    safe_name = _safe_filename(law_id)

    log_export(user_id, ip, law_id, "txt", tier)

    response = HttpResponse(content, content_type="text/plain; charset=utf-8")
    response["Content-Disposition"] = f'attachment; filename="{safe_name}.txt"'
    response["Cache-Control"] = "public, max-age=3600"
    return response


@extend_schema(
    tags=["Export"],
    summary="Export law as PDF",
    description="Download a law's full text as a formatted PDF. Requires a free account.",
)
@api_view(["GET"])
def export_pdf(request, law_id):
    """Export a law as a formatted PDF using WeasyPrint."""
    tier, user_id, ip, error = _check_access(request, "pdf")
    if error:
        return error

    if not _has_weasyprint:
        return Response(
            {"error": "PDF export is not available. WeasyPrint is not installed."},
            status=501,
        )

    law = get_object_or_404(Law, official_id=law_id)
    articles = _get_articles(law_id)

    if not articles:
        return Response({"error": "No articles found for this law."}, status=404)

    ctx = _law_context(law, articles)
    html_string = render_to_string("export/law_pdf.html", ctx)
    pdf_bytes = WeasyHTML(string=html_string).write_pdf()

    safe_name = _safe_filename(law_id)
    log_export(user_id, ip, law_id, "pdf", tier)

    response = HttpResponse(pdf_bytes, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{safe_name}.pdf"'
    response["Cache-Control"] = "public, max-age=3600"
    return response


@extend_schema(
    tags=["Export"],
    summary="Export law as LaTeX",
    description="Download a law as a compilable .tex document. Requires a premium account.",
)
@api_view(["GET"])
def export_latex(request, law_id):
    """Export a law as a LaTeX (.tex) file using Jinja2."""
    tier, user_id, ip, error = _check_access(request, "latex")
    if error:
        return error

    if not _has_jinja2:
        return Response(
            {"error": "LaTeX export is not available. Jinja2 is not installed."},
            status=501,
        )

    law = get_object_or_404(Law, official_id=law_id)
    articles = _get_articles(law_id)

    if not articles:
        return Response({"error": "No articles found for this law."}, status=404)

    import os

    template_dir = os.path.join(os.path.dirname(__file__), "templates", "export")
    env = jinja2.Environment(
        loader=jinja2.FileSystemLoader(template_dir),
        block_start_string=r"\BLOCK{",
        block_end_string="}",
        variable_start_string=r"\VAR{",
        variable_end_string="}",
        comment_start_string=r"\#{",
        comment_end_string="}",
        autoescape=False,
    )

    def latex_escape(s: str) -> str:
        """Escape special LaTeX characters."""
        if not s:
            return ""
        replacements = [
            ("\\", r"\textbackslash{}"),
            ("&", r"\&"),
            ("%", r"\%"),
            ("$", r"\$"),
            ("#", r"\#"),
            ("_", r"\_"),
            ("{", r"\{"),
            ("}", r"\}"),
            ("~", r"\textasciitilde{}"),
            ("^", r"\textasciicircum{}"),
        ]
        for old, new in replacements:
            s = s.replace(old, new)
        return s

    env.filters["latex_escape"] = latex_escape

    ctx = _law_context(law, articles)
    template = env.get_template("law_latex.tex")
    tex_content = template.render(**ctx)

    safe_name = _safe_filename(law_id)
    log_export(user_id, ip, law_id, "latex", tier)

    response = HttpResponse(
        tex_content, content_type="application/x-tex; charset=utf-8"
    )
    response["Content-Disposition"] = f'attachment; filename="{safe_name}.tex"'
    response["Cache-Control"] = "public, max-age=3600"
    return response


@extend_schema(
    tags=["Export"],
    summary="Export law as DOCX",
    description="Download a law as a Word (.docx) document. Requires a premium account.",
)
@api_view(["GET"])
def export_docx(request, law_id):
    """Export a law as a Word (.docx) document using python-docx."""
    tier, user_id, ip, error = _check_access(request, "docx")
    if error:
        return error

    if not _has_docx:
        return Response(
            {"error": "DOCX export is not available. python-docx is not installed."},
            status=501,
        )

    law = get_object_or_404(Law, official_id=law_id)
    articles = _get_articles(law_id)

    if not articles:
        return Response({"error": "No articles found for this law."}, status=404)

    doc = DocxDocument()

    # Title
    title = doc.add_heading(law.name, level=0)
    title.alignment = 1  # center

    # Metadata
    tier_label = _tier_label(law.tier)
    meta_parts = [f"Tipo: {tier_label}"]
    if law.category:
        meta_parts.append(f"Categoría: {law.category}")
    if law.status:
        meta_parts.append(f"Estado: {law.status}")
    meta_parts.append(f"Artículos: {len(articles)}")

    latest_version = law.versions.order_by("-publication_date").first()
    if latest_version and latest_version.publication_date:
        meta_parts.append(f"Publicado: {latest_version.publication_date}")

    meta_para = doc.add_paragraph(" | ".join(meta_parts))
    meta_para.alignment = 1

    doc.add_paragraph("")  # spacer

    # Articles
    for article in articles:
        doc.add_heading(article["article_id"], level=2)
        doc.add_paragraph(article["text"])

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph(
        f"Generado por Tezca — El Espejo de la Ley | {datetime.now().strftime('%Y-%m-%d')} | tezca.mx"
    )
    footer.alignment = 1
    for run in footer.runs:
        run.font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = _safe_filename(law_id)
    log_export(user_id, ip, law_id, "docx", tier)

    response = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{safe_name}.docx"'
    response["Cache-Control"] = "public, max-age=3600"
    return response


@extend_schema(
    tags=["Export"],
    summary="Export law as EPUB",
    description="Download a law as an EPUB e-book. Requires a premium account.",
)
@api_view(["GET"])
def export_epub(request, law_id):
    """Export a law as an EPUB e-book using ebooklib."""
    tier, user_id, ip, error = _check_access(request, "epub")
    if error:
        return error

    if not _has_ebooklib:
        return Response(
            {"error": "EPUB export is not available. ebooklib is not installed."},
            status=501,
        )

    law = get_object_or_404(Law, official_id=law_id)
    articles = _get_articles(law_id)

    if not articles:
        return Response({"error": "No articles found for this law."}, status=404)

    book = epub.EpubBook()
    book.set_identifier(f"tezca-{law_id}")
    book.set_title(law.name)
    book.set_language("es")
    book.add_author("Tezca — El Espejo de la Ley")

    style = epub.EpubItem(
        uid="style",
        file_name="style/default.css",
        media_type="text/css",
        content=b"body { font-family: serif; line-height: 1.6; } "
        b"h1 { text-align: center; } "
        b"h2 { color: #2a2a6e; margin-top: 1.5em; } "
        b".meta { text-align: center; color: #666; font-size: 0.9em; } "
        b".footer { text-align: center; color: #999; font-size: 0.8em; margin-top: 2em; }",
    )
    book.add_item(style)

    # Cover chapter
    tier_label = _tier_label(law.tier)
    cover_html = f"<h1>{_epub_escape(law.name)}</h1>"
    cover_html += f'<p class="meta">Tipo: {_epub_escape(tier_label)}'
    if law.category:
        cover_html += f" | Categoría: {_epub_escape(law.category)}"
    cover_html += f" | {len(articles)} artículos</p>"

    cover = epub.EpubHtml(title="Portada", file_name="cover.xhtml", lang="es")
    cover.content = cover_html
    cover.add_item(style)
    book.add_item(cover)

    # Split articles into chapters (~50 per chapter)
    chapters = []
    chunk_size = 50
    for i in range(0, len(articles), chunk_size):
        chunk = articles[i : i + chunk_size]
        first_id = chunk[0]["article_id"]
        last_id = chunk[-1]["article_id"]
        ch_title = f"{first_id} — {last_id}" if len(chunk) > 1 else first_id

        html_parts = [f"<h1>{_epub_escape(ch_title)}</h1>"]
        for art in chunk:
            html_parts.append(
                f"<h2>{_epub_escape(art['article_id'])}</h2>"
                f"<p>{_epub_escape(art['text'])}</p>"
            )

        ch = epub.EpubHtml(
            title=ch_title,
            file_name=f"chapter_{i // chunk_size + 1}.xhtml",
            lang="es",
        )
        ch.content = "\n".join(html_parts)
        ch.add_item(style)
        book.add_item(ch)
        chapters.append(ch)

    # TOC and spine
    book.toc = [cover] + chapters
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = ["nav", cover] + chapters

    buf = io.BytesIO()
    epub.write_epub(buf, book, {})
    buf.seek(0)

    safe_name = _safe_filename(law_id)
    log_export(user_id, ip, law_id, "epub", tier)

    response = HttpResponse(buf.getvalue(), content_type="application/epub+zip")
    response["Content-Disposition"] = f'attachment; filename="{safe_name}.epub"'
    response["Cache-Control"] = "public, max-age=3600"
    return response


def _epub_escape(s: str) -> str:
    """Escape HTML entities for EPUB XHTML content."""
    if not s:
        return ""
    return (
        s.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


@extend_schema(
    tags=["Export"],
    summary="Export law as JSON",
    description="Download a law's full metadata and articles as structured JSON. Requires a premium account.",
)
@api_view(["GET"])
def export_json(request, law_id):
    """Export a law as structured JSON with metadata + articles."""
    tier, user_id, ip, error = _check_access(request, "json")
    if error:
        return error

    law = get_object_or_404(Law, official_id=law_id)
    articles = _get_articles(law_id)

    if not articles:
        return Response({"error": "No articles found for this law."}, status=404)

    latest_version = law.versions.order_by("-publication_date").first()
    pub_date = None
    if latest_version and latest_version.publication_date:
        pub_date = str(latest_version.publication_date)

    data = {
        "meta": {
            "official_id": law.official_id,
            "name": law.name,
            "short_name": law.short_name,
            "tier": law.tier,
            "category": law.category,
            "state": law.state,
            "status": law.status,
            "law_type": law.law_type,
            "publication_date": pub_date,
            "source_url": law.source_url,
            "article_count": len(articles),
            "exported_at": datetime.now().isoformat(),
            "source": "Tezca — El Espejo de la Ley | tezca.mx",
        },
        "articles": articles,
    }

    json_str = json_module.dumps(data, ensure_ascii=False, indent=2)

    safe_name = _safe_filename(law_id)
    log_export(user_id, ip, law_id, "json", tier)

    response = HttpResponse(json_str, content_type="application/json; charset=utf-8")
    response["Content-Disposition"] = f'attachment; filename="{safe_name}.json"'
    response["Cache-Control"] = "public, max-age=3600"
    return response


@extend_schema(
    tags=["Export"],
    summary="Get export quota info",
    description="Returns the user's current tier, usage, limit, and available formats.",
)
@api_view(["GET"])
def export_quota(request, law_id):
    """Return quota information for the current user/IP."""
    tier, user_id = _get_user_tier(request)
    ip = _get_client_ip(request)

    from .export_throttles import get_export_count

    if user_id:
        used = get_export_count(user_id=user_id)
    else:
        used = get_export_count(ip_address=ip)

    limit = TIER_LIMITS.get(tier, 10)

    formats_available = [
        fmt
        for fmt, req_tier in FORMAT_TIERS.items()
        if TIER_RANK.get(tier, 0) >= TIER_RANK.get(req_tier, 0)
    ]

    return Response(
        {
            "tier": tier,
            "used": used,
            "limit": limit,
            "remaining": max(limit - used, 0),
            "formats_available": formats_available,
        }
    )
